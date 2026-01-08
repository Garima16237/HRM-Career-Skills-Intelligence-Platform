import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import tempfile
from io import BytesIO
from docx import Document
from docx.shared import Inches
from agent import career_agent

# ------------------ PAGE CONFIG ------------------
st.set_page_config(page_title="Career Intelligence Agent", layout="wide")

# ------------------ HEADER WITH LOGO ------------------
col1, col2 = st.columns([1, 8])
with col1:
    st.image(
        "https://www.gartner.com/pi/vendorimages/compunnel_1697620403838.png",
        width=90
    )
with col2:
    st.markdown("""
    <h1>Career & Skills Intelligence Platform</h1>
    <h4 style='color:gray'>Enterprise Career Agent • HR & Manager Decisions</h4>
    """, unsafe_allow_html=True)
st.divider()

# ------------------ SIDEBAR ------------------
st.sidebar.header("Analysis Configuration")
analysis_mode = st.sidebar.radio(
    "Analysis Mode",
    ["Employee Record Analysis", "Scenario / What-if Analysis"]
)
view_mode = st.sidebar.selectbox("View Mode", ["Manager", "HR"])

uploaded_file = None
df = None
if analysis_mode == "Employee Record Analysis":
    uploaded_file = st.sidebar.file_uploader("Upload Employee CSV", type="csv")
    if uploaded_file:
        df = pd.read_csv(uploaded_file)

# ------------------ EMPLOYEE DATA ------------------
emp = {}
if analysis_mode == "Employee Record Analysis" and df is not None:
    selected_emp_id = st.sidebar.selectbox(
        "Select Employee ID",
        df["employee_id"].astype(str).unique()
    )
    emp = df[df["employee_id"].astype(str) == selected_emp_id].iloc[0]

# ------------------ INPUT DATA ------------------
if analysis_mode == "Employee Record Analysis":
    # read-only
    employee_id = emp.get("employee_id", "")
    name = emp.get("name", "")
    role = emp.get("role", "")
    target_role = emp.get("target_role", "")
    experience = int(emp.get("experience", 0))
    skills = emp.get("skills", "")
    certifications = emp.get("certifications", "")
else:
    st.sidebar.subheader("Scenario Inputs")
    employee_id = st.sidebar.text_input("Employee ID (optional)")
    name = st.sidebar.text_input("Employee Name")
    role = st.sidebar.text_input("Current Role")
    target_role = st.sidebar.text_input("Target Role")
    experience = st.sidebar.number_input("Experience (years)", 0, 40)
    skills = st.sidebar.text_area("Skills (comma separated)")
    certifications = st.sidebar.text_area("Certifications (if any)")

skill_list = [s.strip() for s in skills.split(",") if s.strip()]

# ------------------ SELF-ASSESSMENT ------------------
st.sidebar.subheader("Self-Assessment")
self_confidence = st.sidebar.selectbox(
    "Overall Skill Confidence",
    ["Foundational", "Working Professional", "Advanced", "Expert"]
)
ownership_level = st.sidebar.selectbox(
    "Primary Responsibility Level",
    [
        "Execution-focused",
        "Independent contributor",
        "Module owner",
        "End-to-end owner"
    ]
)

# ------------------ HR APPROVAL ------------------
approval_status = "Draft"
if view_mode == "HR":
    approval_status = st.sidebar.selectbox(
        "Approval Status",
        ["Draft", "Approved", "Rejected"]
    )

# ------------------ SCORING ------------------
career_readiness = min(60 + experience*3, 90)
promotion_score = int(career_readiness*0.85)
peer_percentile = max(55, min(95, 60 + (career_readiness-55)))

# ------------------ GROQ AGENT PROMPT ------------------
def generate_insights():
    prompt = f"""
You are a Senior Enterprise HR Career Intelligence Agent.

Employee ID: {employee_id}
Name: {name}
Current Role: {role}
Target Role: {target_role}
Experience: {experience} years
Skills: {skills}
Certifications: {certifications}

Self-Assessment:
Confidence: {self_confidence}
Responsibility: {ownership_level}

Rules:
- No numeric skill ratings
- Executive HR-safe language
- Strategic promotion framing

Deliver:
Executive overview, skills assessment,
career path, readiness statement,
roadmap, HR notes.
"""
    return career_agent(prompt)

# ------------------ VISUALS ------------------
def create_skill_plot():
    fig, ax = plt.subplots(figsize=(6,3))
    core_skills = ["Python","ML","AI","SQL","NLP","Azure","PowerBI","Leadership"]
    available = set(skill_list)
    covered = [s for s in core_skills if s in available]
    missing = [s for s in core_skills if s not in available]
    ax.barh(covered, [1]*len(covered), color="green", label="Aligned Skills")
    ax.barh(missing, [1]*len(missing), left=[0]*len(missing), color="red", label="Development Area")
    ax.set_xticks([])
    ax.set_title("Skill Coverage vs Gaps")
    ax.legend()
    return fig

def create_readiness_plot():
    fig, ax = plt.subplots(figsize=(4,3))
    ax.bar(["Career Readiness","Promotion"], [career_readiness, promotion_score], color=["blue","orange"])
    ax.set_ylim(0,100)
    ax.set_title("Career Positioning Overview")
    return fig

# ------------------ MAIN PANELS ------------------
left, right = st.columns([1.3,2.7])

with left:
    if analysis_mode == "Employee Record Analysis":
        st.subheader("👤 Employee Snapshot")
        st.text_input("Employee ID", employee_id, disabled=True)
        st.text_input("Name", name, disabled=True)
        st.text_input("Current Role", role, disabled=True)
        st.text_input("Target Role", target_role, disabled=True)
        st.text_input("Experience", f"{experience} years", disabled=True)
    else:
        st.subheader("✏️ Scenario Inputs")
        st.info("Editable for what-if analysis")

with right:
    st.subheader("🧠 Career Intelligence Output")
    output_container = st.container()

    if st.button("Run Career Analysis"):
        with st.spinner("Generating insights..."):
            insights = generate_insights()
            st.session_state["insights"] = insights
            output_container.markdown(insights)

            # Table
            st.markdown("### 📋 Skills Overview")
            skill_table = pd.DataFrame({
                "Skill": skill_list,
                "Status": ["Aligned" if s in skill_list else "Development Needed" for s in skill_list],
                "Recommendation": ["Keep/Develop"]*len(skill_list)
            })
            st.table(skill_table)

            # Charts
            st.markdown("### 📊 Career Charts")
            st.pyplot(create_skill_plot())
            st.pyplot(create_readiness_plot())

# ------------------ DOCX DOWNLOAD ------------------
if approval_status=="Approved" and "insights" in st.session_state:
    st.success("HR Approved — DOC download enabled")
    if st.button("Download Career Report (DOCX)"):
        doc = Document()
        doc.add_heading("Career Intelligence Report", 0)
        doc.add_paragraph(st.session_state["insights"])

        # Save charts as images
        skill_fig = create_skill_plot()
        readiness_fig = create_readiness_plot()
        skill_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        readiness_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        skill_fig.savefig(skill_img.name)
        readiness_fig.savefig(readiness_img.name)

        doc.add_heading("Skill Coverage vs Gaps", level=1)
        doc.add_picture(skill_img.name, width=Inches(5))
        doc.add_heading("Career Positioning Overview", level=1)
        doc.add_picture(readiness_img.name, width=Inches(5))

        tmp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp_doc.name)
        with open(tmp_doc.name,"rb") as f:
            st.download_button(
                "Download Full Career Report (DOCX)",
                f,
                file_name=f"{employee_id}_Career_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

st.info("Enterprise Career Intelligence Agent • Compunnel")
